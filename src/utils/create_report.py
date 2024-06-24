import os
import sys
from docx import Document
from docx.shared import Mm, Inches, Pt


def create_report_avd(imgdir, output_filename='report.docx'):
    """
    Create a Word document report containing acceleration, velocity, and
    displacement time history plots located in an input directory
    """
    if not os.path.isdir(imgdir):
        raise NotADirectoryError("Provided input is not a valid directory!")

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.italic = True
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)

    style = doc.styles['Header']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(16)
    font.bold = True

    imgpath_list = [f for f in os.listdir(imgdir)
        if f.endswith('.svg') or f.endswith('.png') or f.endswith('.jpg')]

    num_set = set()
    comp_set = set()
    im_set = set()
    for img in imgpath_list:
        iname, _ = img.split('.')
        num, comp, im = iname.split()
        if im.upper() in 'ACC' or im.upper() in 'VEL' or im.upper() in 'DISP':
            num_set.add(num)
            comp_set.add(comp)
            im_set.add(im)

    num_list = list(num_set)
    num_list.sort()
    comp_list = list(comp_set)
    comp_list.sort()
    im_list = list(im_set)
    im_list.sort()

    for num in num_list:
        for comp in comp_list:
            new_section = doc.add_section()

            new_section.page_height = Mm(297)
            new_section.page_width = Mm(210)
            new_section.left_margin = new_section.right_margin = \
                new_section.top_margin = new_section.bottom_margin = Inches(1)
            acc, disp, vel = im_list

            _ = doc.add_paragraph('Acceleration')
            doc.add_picture(
                os.path.join(imgdir, ' '.join([num, comp, acc]) + '.png'),
                width=Inches(6.27))

            _ = doc.add_paragraph()
            _ = doc.add_paragraph('Velocity')
            doc.add_picture(
                os.path.join(imgdir, ' '.join([num, comp, vel]) + '.png'),
                width=Inches(6.27))
            
            _ = doc.add_paragraph()
            _ = doc.add_paragraph('Displacement')
            doc.add_picture(
                os.path.join(imgdir, ' '.join([num, comp, disp]) + '.png'),
                width=Inches(6.27))
            
            header = new_section.header
            header.is_linked_to_previous = False
            header_paragraph = header.paragraphs[0]
            
            header_run1 = header_paragraph.add_run(f"[{num}] ")
            header_run1.italic = False
            header_run2 = header_paragraph.add_run(f"{comp} Component")
            header_run2.italic = False
            header_run2.underline = True

    doc.save(os.path.join(imgdir, output_filename))
    print("Report generated successfully...")

if __name__ == '__main__':
    images_directory = sys.argv[1]
    create_report_avd(images_directory)